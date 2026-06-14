from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path(r"d:\LMSPROJECT\artifacts\lms-qa-testing-report.docx")

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x0B, 0x25, 0x45)
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_BLUE_FILL = "E8EEF5"
LIGHT_GRAY_FILL = "F2F4F7"
GRID = "C9D2DE"


def set_font(run, name, size_pt, color=None, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Calibri"
    heading1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading1.font.size = Pt(16)
    heading1.font.color.rgb = BLUE
    heading1.font.bold = True
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(10)
    heading1.paragraph_format.line_spacing = 1.25

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Calibri"
    heading2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading2.font.size = Pt(13)
    heading2.font.color.rgb = BLUE
    heading2.font.bold = True
    heading2.paragraph_format.space_before = Pt(14)
    heading2.paragraph_format.space_after = Pt(7)
    heading2.paragraph_format.line_spacing = 1.25

    heading3 = doc.styles["Heading 3"]
    heading3.font.name = "Calibri"
    heading3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading3.font.size = Pt(12)
    heading3.font.color.rgb = DARK_BLUE
    heading3.font.bold = True
    heading3.paragraph_format.space_before = Pt(10)
    heading3.paragraph_format.space_after = Pt(5)
    heading3.paragraph_format.line_spacing = 1.25

    list_bullet = doc.styles["List Bullet"]
    list_bullet.font.name = "Calibri"
    list_bullet._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    list_bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    list_bullet.font.size = Pt(11)
    list_bullet.paragraph_format.space_after = Pt(4)
    list_bullet.paragraph_format.line_spacing = 1.25
    list_bullet.paragraph_format.left_indent = Inches(0.375)
    list_bullet.paragraph_format.first_line_indent = Inches(-0.188)

    list_number = doc.styles["List Number"]
    list_number.font.name = "Calibri"
    list_number._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    list_number._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    list_number.font.size = Pt(11)
    list_number.paragraph_format.space_after = Pt(4)
    list_number.paragraph_format.line_spacing = 1.25
    list_number.paragraph_format.left_indent = Inches(0.375)
    list_number.paragraph_format.first_line_indent = Inches(-0.188)

    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run("LMS QA Testing Review")
    set_font(run, "Calibri", 9, MUTED)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Learning Management System (LMS)\nQA Testing Review Report")
    set_font(run, "Calibri", 22, INK, bold=True)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after = Pt(10)
    run = sub.add_run(
        f"Prepared for system testing and professional QA review\nGenerated on {date.today().strftime('%B %d, %Y')}"
    )
    set_font(run, "Calibri", 11, MUTED)


def add_summary_table(doc):
    doc.add_heading("Executive Summary", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    widths = [1.9, 4.6]
    labels_and_values = [
        ("Application Scope", "Auth, dashboards, user management, courses, batches, video content, quizzes, chatbot, and doubt-session scheduling."),
        ("Current Test Baseline", "Existing automated suite passes with 86 tests and 3 skipped checks."),
        ("Highest Risks", "Role-based access leaks, upload abuse, quiz workflow regressions, and client-side XSS in shared chatbot/notification rendering."),
        ("Major Gaps", "Assignments, certificates, full reporting, persistent progress tracking, and true in-app notifications are not evident as completed modules."),
        ("Recommended First Pass", "Run role-based end-to-end flows for Admin, Trainer, Student, and Guest before deeper security and performance testing."),
    ]

    for row, (label, value) in zip(table.rows, labels_and_values):
        for idx, width in enumerate(widths):
            set_cell_width(row.cells[idx], width)
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(row.cells[0], LIGHT_BLUE_FILL)
        p0 = row.cells[0].paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(label)
        set_font(r0, "Calibri", 11, DARK_BLUE, bold=True)

        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(value)
        set_font(r1, "Calibri", 11, INK)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_font(r, "Calibri", 11, INK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_font(r, "Calibri", 11, INK)


def add_section(doc, title, intro, bullets):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph(intro)
    p.paragraph_format.space_after = Pt(6)
    add_bullets(doc, bullets)


def add_bug_table(doc):
    doc.add_heading("Known Issues Already Identified", level=1)
    intro = doc.add_paragraph(
        "The following issues are already visible from the current repository review and should be included in the first QA execution cycle."
    )
    intro.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    headers = ["Bug Title", "Expected Result", "Actual Result", "Severity / Priority"]
    widths = [2.0, 1.7, 1.9, 0.9]

    for i, (header, width) in enumerate(zip(headers, widths)):
        set_cell_width(table.rows[0].cells[i], width)
        set_cell_shading(table.rows[0].cells[i], LIGHT_GRAY_FILL)
        p = table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(header)
        set_font(r, "Calibri", 10, DARK_BLUE, bold=True)

    issues = [
        (
            "Chatbot / toast rendering allows unsafe HTML sinks",
            "User and system text should render as plain safe content.",
            "Shared layout uses innerHTML and insertAdjacentHTML in multiple places.",
            "High / High",
        ),
        (
            "Assignments, certificates, grading, and reports not available",
            "Core LMS modules should exist if listed in product scope.",
            "Repo scan shows no complete assignment, certificate, or grading/report modules.",
            "High / Medium",
        ),
        (
            "Notification bell is placeholder-only",
            "Notifications should persist and show actual events.",
            "Topbar menu shows a static empty message and client-side toast helper only.",
            "Medium / Medium",
        ),
        (
            "Persistent lesson progress tracking unclear",
            "Lesson/video progress should be saved and visible after refresh and re-login.",
            "Quiz attempts exist, but lesson-completion tracking is not evident.",
            "Medium / High",
        ),
    ]

    for issue in issues:
        row = table.add_row()
        for idx, (value, width) in enumerate(zip(issue, widths)):
            set_cell_width(row.cells[idx], width)
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = row.cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(value)
            set_font(r, "Calibri", 10, INK)


def add_bug_report_template(doc):
    doc.add_heading("Bug Reporting Template", level=1)
    add_numbered(doc, [
        "Bug title",
        "Module and test environment",
        "Preconditions",
        "Steps to reproduce",
        "Expected result",
        "Actual result",
        "Severity",
        "Priority",
        "Screenshot or video suggestion",
        "Recommended fix",
    ])


def add_final_summary(doc):
    doc.add_heading("Final QA Report Structure", level=1)
    add_bullets(doc, [
        "Major findings by module and role.",
        "Critical security issues and authorization risks.",
        "UI and usability inconsistencies.",
        "Performance observations with evidence.",
        "Data-integrity and regression concerns.",
        "Open product-scope gaps and enhancement recommendations.",
    ])

    doc.add_heading("Recommended Execution Order", level=2)
    add_numbered(doc, [
        "Smoke test login, logout, registration, and dashboard access for all roles.",
        "Run the main admin-to-trainer-to-student course and quiz workflow.",
        "Probe direct URL access and role-based authorization boundaries.",
        "Stress file uploads, quiz submissions, and chatbot queries with invalid and edge inputs.",
        "Execute responsive, cross-browser, and mobile checks on high-traffic screens.",
        "Perform targeted security testing for XSS, CSRF, session handling, and upload abuse.",
        "Finish with regression confirmation after fixes.",
    ])


def main():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_styles(doc)
    add_title(doc)
    add_summary_table(doc)

    doc.add_heading("Application Scope Under Test", level=1)
    add_bullets(doc, [
        "Implemented modules visible in the current project: authentication, admin/instructor/student dashboards, user management, courses, batch assignment, videos/study content, quizzes, chatbot, and doubt sessions.",
        "Modules requested in a generic LMS checklist but not clearly implemented here: assignments, certificates, gradebook/grading, real reports, and robust notification management.",
        "QA execution should distinguish between true defects and product-scope gaps.",
    ])

    add_section(
        doc,
        "Functional Testing",
        "Validate each implemented feature against the expected LMS workflow and confirm that page actions save correct state.",
        [
            "Verify guest access to login and registration, then confirm blocked access to protected pages.",
            "Verify role-specific redirects after login for Admin, Trainer, and Student.",
            "Confirm admin user management: create, activate/deactivate, safe delete, and search/filter flows.",
            "Confirm course creation, editing, trainer assignment, trainee assignment, batch creation, and batch-based enrollment.",
            "Verify instructor content flow: upload video, upload study material, open content detail, and quiz-generation workflow.",
            "Verify student content flow: browse course, access enrolled content only, take quiz once, and view results.",
            "Verify doubt-session request, slot proposal, slot selection, confirmation, postponement, and closure flows.",
        ],
    )

    add_section(
        doc,
        "UI / UX and Consistency Testing",
        "Check whether the LMS feels polished, consistent, and professional across core screens.",
        [
            "Review page titles, labels, alert messages, spelling, and role terminology consistency.",
            "Check sidebar, topbar, buttons, cards, table spacing, and icon alignment across roles.",
            "Verify empty states, validation states, success messages, and destructive-action confirmations.",
            "Check responsive behavior for login pages, dashboards, course tables, forms, and mobile sidebar interactions.",
            "Validate font consistency, color hierarchy, and whether actions remain obvious without visual clutter.",
        ],
    )

    add_section(
        doc,
        "Role-Based Access Testing",
        "Systematically verify that each role sees only its allowed data, pages, and actions.",
        [
            "Guest: direct URL access to dashboards, courses, videos, quizzes, and session routes should redirect to login.",
            "Student: confirm no access to admin user management, course editing, quiz review, or instructor session actions.",
            "Trainer: confirm access only to assigned courses and own review/session workflows.",
            "Admin: confirm unrestricted management access while preventing self-destructive account actions where intended.",
            "Inactive users: confirm login and access behavior is appropriate after admin deactivation.",
        ],
    )

    add_section(
        doc,
        "Security Testing",
        "Use a practical web-app security pass focused on the most realistic LMS risks.",
        [
            "Authentication: brute-force login lockout, logout behavior, browser back-button after logout, and multiple-session behavior.",
            "Authorization: direct URL manipulation, object-ID tampering, and cross-role POST attempts.",
            "XSS: test chatbot input, quiz text, profile fields, course descriptions, and any rendered rich text or notifications.",
            "CSRF: submit protected forms and AJAX endpoints without tokens and confirm rejection.",
            "File uploads: wrong file type, renamed executable, corrupted image, oversized file, and interrupted upload attempts.",
            "Error handling: verify no stack traces, secrets, or internal implementation details leak to users.",
            "Sensitive data: inspect HTML, JavaScript, network responses, and logs for tokens or hidden data exposure.",
        ],
    )

    add_section(
        doc,
        "Performance Testing",
        "Measure the LMS under realistic usage patterns and slower network conditions.",
        [
            "Time first meaningful render for login, admin dashboard, student dashboard, course list, and video detail.",
            "Measure study-material processing feedback, quiz submission speed, and chatbot response latency.",
            "Observe course lists and dashboard performance with larger user and course datasets.",
            "Run concurrent smoke load on login, dashboard opens, course browsing, and quiz submission endpoints.",
            "Check whether polling or repeated requests cause unnecessary delay or browser lag.",
        ],
    )

    add_section(
        doc,
        "Compatibility Testing",
        "Confirm the LMS behaves consistently across common browser and device combinations.",
        [
            "Desktop browsers: Chrome, Edge, Firefox, and Safari when available.",
            "Screen sizes: small mobile, tablet portrait, tablet landscape, laptop, and wide desktop.",
            "Check touch interactions for menus, dropdowns, tables, and forms on mobile devices.",
            "Verify video/content pages, admin tables, and quiz forms remain usable on narrow screens.",
        ],
    )

    add_section(
        doc,
        "Form and Validation Testing",
        "Exercise every important input path with valid, invalid, duplicate, and boundary data.",
        [
            "Registration and user creation: duplicate usernames, duplicate emails, weak passwords, empty fields, and role selection.",
            "Profile updates: invalid email, bad Meet link, invalid profile picture, and large image uploads.",
            "Course and batch forms: empty titles, duplicate names, and invalid assignment combinations.",
            "Quiz and doubt-session forms: incomplete answers, invalid slot values, malformed POST data, and repeated submissions.",
            "Confirm that validation messages are accurate, readable, and attached to the correct fields.",
        ],
    )

    add_section(
        doc,
        "Course and Exam Testing",
        "Focus on the admin-instructor-student academic workflow end to end.",
        [
            "Verify course creation, trainer assignment, trainee enrollment, and batch-driven enrollment.",
            "Confirm student access is blocked until enrollment is active.",
            "Verify study material upload prerequisites for quiz generation.",
            "Confirm draft review, approval, rejection, replacement generation, publishing, and student visibility.",
            "Verify quiz scoring, result display, answer integrity, and retake restriction behavior.",
        ],
    )

    add_section(
        doc,
        "Data Integrity Testing",
        "Confirm that the LMS stores the correct state and does not lose it unexpectedly.",
        [
            "Check persistence of enrollments, active/inactive user state, quiz attempts, and session states after refresh and re-login.",
            "Verify that course assignments and batch assignments remain consistent after edits.",
            "Confirm score calculations and saved answers remain unchanged after navigation or logout.",
            "Check whether deactivation and safe-delete flows preserve dependent records as intended.",
        ],
    )

    add_section(
        doc,
        "Negative and Edge Case Testing",
        "Use failure-oriented testing to expose hidden bugs and workflow weaknesses.",
        [
            "Double-submit forms and quiz actions to look for duplicate writes.",
            "Refresh or navigate back during uploads, quiz attempts, and session-selection flows.",
            "Try expired or invalid IDs in direct URLs for courses, videos, quiz drafts, and session actions.",
            "Open the LMS in multiple tabs and test for stale UI or conflicting state.",
            "Test inactive users, missing related records, and partial workflows interrupted midstream.",
        ],
    )

    add_bug_table(doc)
    add_bug_report_template(doc)
    add_final_summary(doc)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
